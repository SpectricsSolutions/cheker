# import streamlit as st
# import pandas as pd
# from docx import Document
# from collections import Counter

# st.set_page_config(page_title="CSV Completeness Checker", layout="wide")

# # === Load Rules from Word DOCX ===
# def load_rules(docx_file):
#     doc = Document(docx_file)
#     rules = {}
#     current_type = None
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if text.lower() in ["annual review", "pension transfer", "new business", "ad hoc withdrawal"]:
#             current_type = text.strip().lower()
#             rules[current_type] = []
#         elif current_type and text:
#             rules[current_type].append(text.strip().lower())
#     return rules

# # === Analyze CSV Rows ===
# def analyze_csv(csv_df, rules):
#     results = []
#     issues_fields = Counter()

#     for _, row in csv_df.iterrows():
#         report_type = str(row.get("report type", "")).strip().lower()
#         required_fields = rules.get(report_type, [])

#         missing_fields = []
#         for field in required_fields:
#             value = row.get(field.lower(), "")
#             if pd.isna(value) or str(value).strip().lower() in ["", "none", "nan"]:
#                 missing_fields.append(field)
#                 issues_fields[field] += 1

#         if not missing_fields:
#             status = "complete"
#         elif len(missing_fields) <= 2:
#             status = "incomplete"
#         else:
#             status = "missing data"

#         results.append({
#             "client": row.get("client name", "Unknown"),
#             "report_type": report_type,
#             "status": status,
#             "missing": missing_fields
#         })

#     return results, issues_fields

# # === Streamlit UI ===
# st.title("📊 CSV Completeness Checker (by Report Type)")

# csv_file = st.file_uploader("📥 Upload CSV File", type=["csv"])
# rules_docx = st.file_uploader("📘 Upload Rules (DOCX)", type=["docx"])

# if csv_file and rules_docx:
#     df = pd.read_csv(csv_file)
#     df.columns = df.columns.str.strip().str.lower()

#     rules = load_rules(rules_docx)
#     results, issues_fields = analyze_csv(df, rules)

#     total = len(results)
#     complete = sum(1 for r in results if r["status"] == "complete")
#     incomplete = total - complete

#     st.subheader("📈 Summary")
#     col1, col2, col3 = st.columns(3)
#     col1.metric("Total Cases", total)
#     col2.metric("✅ Complete", complete)
#     col3.metric("❌ Incomplete", incomplete)

#     st.subheader("🧾 Breakdown by Report Type")
#     df_res = pd.DataFrame(results)
#     if not df_res.empty:
#         breakdown = df_res.groupby(['report_type', 'status']).size().unstack(fill_value=0)
#         st.dataframe(breakdown)

#     st.subheader("⚠️ Common Missing Fields")
#     if issues_fields:
#         st.dataframe(pd.DataFrame(issues_fields.items(), columns=["Field", "Missing Count"]).sort_values(by="Missing Count", ascending=False))
#     else:
#         st.info("No missing fields found.")

#     st.subheader("📋 Detailed Results")
#     for r in results:
#         with st.expander(f"{r['client']} ({r['report_type'].title()}): {r['status'].upper()}"):
#             if r["missing"]:
#                 st.write("Missing Fields:", r["missing"])
#             else:
#                 st.write("All required fields present ✅")

import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from collections import Counter

st.set_page_config(page_title="📋 Data & Document Completeness Checker", layout="wide")

# === Load Rules from DOCX ===
def load_rules(docx_file):
    doc = Document(docx_file)
    rules = {}
    current_type = None
    mode = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower() in ["annual review", "pension transfer", "new business", "ad hoc withdrawal"]:
            current_type = text.lower()
            rules[current_type] = {"fields": [], "documents": []}
        elif text.lower() == "required data fields:":
            mode = "fields"
        elif text.lower() == "required documents:":
            mode = "documents"
        elif current_type and mode and text:
            rules[current_type][mode].append(text.strip().lower())

    return rules

# === Extract text from PDFs ===
def extract_text_from_pdf(file):
    text = ""
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        st.warning(f"⚠️ Could not read PDF {file.name}: {e}")
    return text.lower()

# === Analyze CSV & PDFs ===
def analyze_data(csv_df, pdf_files, rules):
    results = []
    field_issues = Counter()
    pdf_texts = {pdf.name: extract_text_from_pdf(pdf) for pdf in pdf_files}
    combined_text = " ".join(pdf_texts.values())

    for _, row in csv_df.iterrows():
        client_name = str(row.get("client name", "")).strip().lower()
        report_type = str(row.get("report type", "")).strip().lower()

        rule = rules.get(report_type, {})
        required_fields = rule.get("fields", [])
        required_docs = rule.get("documents", [])

        # Field check
        missing_fields = []
        for field in required_fields:
            value = row.get(field.lower(), "")
            if pd.isna(value) or str(value).strip().lower() in ["", "none", "nan"]:
                missing_fields.append(field)
                field_issues[field] += 1

        # Document check
        matched_docs = []
        unmatched_docs = []
        for doc_name in required_docs:
            if doc_name.lower() in combined_text and client_name in combined_text:
                matched_docs.append(doc_name)
            else:
                unmatched_docs.append(doc_name)

        status = "✅ Complete"
        if missing_fields or unmatched_docs:
            status = "❌ Incomplete"

        results.append({
            "client": client_name.title() or "Unknown",
            "report_type": report_type.title(),
            "status": status,
            "missing_fields": missing_fields,
            "matched_documents": matched_docs,
            "missing_documents": unmatched_docs
        })

    return results, field_issues

# === Streamlit UI ===
st.title("📋 CSV & PDF Completeness Checker")

# Uploads
csv_file = st.file_uploader("📥 Upload CSV File", type=["csv"])
rules_file = st.file_uploader("📘 Upload Rules DOCX", type=["docx"])
pdf_files = st.file_uploader("📄 Upload One or More PDFs", type=["pdf"], accept_multiple_files=True)

if csv_file and rules_file and pdf_files:
    csv_df = pd.read_csv(csv_file)
    csv_df.columns = csv_df.columns.str.strip().str.lower()

    rules = load_rules(rules_file)
    results, issues_fields = analyze_data(csv_df, pdf_files, rules)

    # Summary
    st.subheader("📊 Summary")
    total = len(results)
    complete = sum(1 for r in results if r["status"] == "✅ Complete")
    incomplete = total - complete

    col1, col2, col3 = st.columns(3)
    col1.metric("Total Cases", total)
    col2.metric("✅ Complete", complete)
    col3.metric("❌ Incomplete", incomplete)

    # Breakdown by Report Type
    st.subheader("📁 Breakdown by Report Type")
    df_summary = pd.DataFrame(results)
    breakdown = df_summary.groupby(['report_type', 'status']).size().unstack(fill_value=0)
    st.dataframe(breakdown)

    # Common Missing Fields
    st.subheader("⚠️ Common Missing Fields (from CSV)")
    if issues_fields:
        field_df = pd.DataFrame(issues_fields.items(), columns=["Field", "Missing Count"])
        st.dataframe(field_df.sort_values(by="Missing Count", ascending=False))
    else:
        st.success("✅ No missing fields detected in CSV.")

    # Per-client report
    st.subheader("📋 Row-by-Row Results")
    for r in results:
        with st.expander(f"{r['client']} ({r['report_type']}): {r['status']}"):
            if r["missing_fields"]:
                st.write("❌ Missing Fields:", r["missing_fields"])
            else:
                st.write("✅ All required fields present.")

            if r["matched_documents"]:
                st.write("✅ Matched Documents:", r["matched_documents"])
            else:
                st.write("⚠️ No matched documents found.")

            if r["missing_documents"]:
                st.write("❌ Missing Documents:", r["missing_documents"])
            else:
                st.write("✅ All required documents present.")
else:
    st.info("Please upload all 3 files: CSV, Rules DOCX, and PDFs to begin.")
